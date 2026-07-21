# -*- coding: utf-8 -*-
"""生成《密码管理系统(passwd-web)维护讲解文档.docx》。
所有代码引用均已与磁盘源码核对（行号可定位）。运行：
  python gen_maintenance_doc.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(__file__), "passwd-web维护讲解文档.docx")

doc = Document()

# ---------- 基础字体（中文） ----------
normal = doc.styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal.font.size = Pt(10.5)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
rfonts.set(qn("w:ascii"), "Microsoft YaHei")
rfonts.set(qn("w:hAnsi"), "Microsoft YaHei")

# 标题配色
def _set_color(run, hexcolor):
    run.font.color.rgb = RGBColor.from_string(hexcolor)

# ---------- 辅助函数 ----------
def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    _set_color(r, "1F4E79")
    return p

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    _set_color(r, "2E75B6")
    return p

def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    _set_color(r, "548235")
    return p

def para(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.add_run(text)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.8)
    rr = run._element.get_or_add_rPr().get_or_add_rFonts()
    rr.set(qn("w:ascii"), "Consolas")
    rr.set(qn("w:hAnsi"), "Consolas")
    rr.set(qn("w:cs"), "Consolas")
    return p

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        rp = hdr[i].paragraphs[0].add_run(htext)
        rp.bold = True
        rp.font.size = Pt(9)
        rp.font.name = "Microsoft YaHei"
        rp._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
        _set_cell_bg(hdr[i], "2E75B6")
        _set_cell_text_color(hdr[i], "FFFFFF")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            rp = cells[i].paragraphs[0].add_run(str(val))
            rp.font.size = Pt(8.8)
            rp.font.name = "Microsoft YaHei"
            rp._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t

def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def _set_cell_text_color(cell, hexcolor):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor.from_string(hexcolor)

def note(text, kind="warn"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(6)
    tag = "⚠ 注意" if kind == "warn" else "💡 提示"
    r0 = p.add_run(tag + "：")
    r0.bold = True
    r0.font.size = Pt(9.5)
    _set_color(r0, "C00000" if kind == "warn" else "2E75B6")
    r1 = p.add_run(text)
    r1.font.size = Pt(9.5)
    return p

# ================================================================
# 封面
# ================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("密码管理系统（passwd-web）\n维护讲解文档")
tr.bold = True
tr.font.size = Pt(24)
tr.font.name = "Microsoft YaHei"
tr._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
_set_color(tr, "1F4E79")

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("面向手动维护者的代码级架构与运维说明 v1.0")
sr.font.size = Pt(12)
sr.font.name = "Microsoft YaHei"
sr._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
_set_color(sr, "595959")

doc.add_paragraph()
para("本文档面向需要接手本系统日常维护、二次开发或故障排查的工程人员。内容精细到"
     "关键函数、路由与代码片段，并标注了多租户隔离闸门与常见易踩的坑。"
     "文中所有代码引用均已与仓库磁盘源码核对，括号内行号可供定位。", italic=True)
doc.add_page_break()

# ================================================================
# 1. 目的与读者
# ================================================================
h1("1. 文档目的与读者")
bullet("日常维护者：理解系统如何在离线内网运行、升级时如何不丢数据、出问题时去哪改。")
bullet("二次开发者：知道新增一个密码算法、一个接口、一个管理页面分别要动哪些文件。")
bullet("故障排查者：快速定位“登录失败 / 看不到密码 / 越权 / 解密报错”的根因位置。")
note("本文与代码同步维护。任何一次涉及认证、加密、多租户隔离、数据模型的改动，都必须同步更新本章节对应的“易踩的坑”。", "warn")

# ================================================================
# 2. 项目概述
# ================================================================
h1("2. 项目概述")
para("passwd-web 是一个离线内网部署的零知识密码保险箱。其核心安全模型是：")
bullet("内层：每条密码用独立的“条目密码（entry_password）”经 SM4-CBC 加密，服务端只存密文，不持有条目密码。")
bullet("外层：gpg / sm2 算法的条目，再额外用组织密钥（OrgKey）或服务端默认密钥做一次非对称加密。")
bullet("认证：登录采用 SCRAM-SM3 挑战应答，密码本身不落网、不落库（库里只存 SM3(password||salt)）。")
bullet("多租户：按“分组（group）”隔离，普通用户只能看自己所属分组的数据；管理员分“超级管理员”和“分组管理员”两种形态。")
para("技术栈：后端 FastAPI + SQLAlchemy + SQLite；前端 Vue3 + Vite（构建产物由后端同源托管）；"
     "部署为单个 Docker 镜像（含预编译前端与后端），离线运行。")

# ================================================================
# 3. 技术栈
# ================================================================
h1("3. 技术栈与运行环境")
table(
    ["层", "技术", "说明 / 版本约束"],
    [
        ["后端框架", "FastAPI", "ASGI；路由用 APIRouter 拆分"],
        ["ORM", "SQLAlchemy 2.0", "模型在 models.py，会话工厂在 db.py"],
        ["数据库", "SQLite", "WAL 模式；落盘 backend/data/*.db"],
        ["对称加密", "SM4-CBC + PBKDF2-SM3", "条目内层加密，纯 Python 自实现（entry_cipher.py）"],
        ["非对称加密", "GPG(pgpy) / SM2(gmssl)", "外层 OrgKey 加密"],
        ["哈希", "SM3（国密）", "登录 verifier 与 proof 计算"],
        ["认证令牌", "JWT (HS256 + jti)", "security.create_token"],
        ["前端", "Vue 3 + Vite", "SFC 组件；构建产物 frontend/dist"],
        ["部署", "Docker 单镜像", "python:3.13-slim 基础；run.py 入口"],
        ["依赖约束", "bcrypt==4.0.1", "passlib 与 ≥4.1 不兼容，必须固定"],
    ],
    widths=[1.2, 2.4, 3.2],
)

# ================================================================
# 4. 目录结构
# ================================================================
h1("4. 目录结构总览")
code_block(
"passwd-web-new/\n"
"├─ backend/\n"
"│  ├─ app/                 # 后端源码（已装入镜像 /app/app）\n"
"│  │  ├─ main.py           # FastAPI 入口、静态托管、安全头\n"
"│  │  ├─ run.py            # 启动入口（含 HTTPS / 明文跳转）\n"
"│  │  ├─ config.py         # 配置与密钥持久化\n"
"│  │  ├─ db.py             # 引擎、会话、通用迁移 _migrate_columns\n"
"│  │  ├─ models.py         # 所有 SQLAlchemy 模型\n"
"│  │  ├─ seed.py           # 首次启动建表/管理员/默认密钥\n"
"│  │  ├─ security.py       # 哈希、JWT、SCRAM-SM3\n"
"│  │  ├─ crypto/           # 加密：entry_cipher / manager / gpg / sm2\n"
"│  │  ├─ core/             # deps.py 多租户闸门、sessions.py\n"
"│  │  └─ routers/          # auth / admin / passwords / keys / history ...\n"
"│  ├─ frontend/            # 前端源码（构建后 dist 拷入镜像）\n"
"│  │  ├─ src/\n"
"│  │  │  ├─ main.js / App.vue\n"
"│  │  │  ├─ store.js       # 全局状态与权限\n"
"│  │  │  ├─ api/           # http.js / auth.js\n"
"│  │  │  ├─ crypto/sm3.js  # 前端 SM3 与 SCRAM proof\n"
"│  │  │  └─ components/    # 各 UI 组件\n"
"│  │  └─ dist/             # 已构建静态产物（打进镜像）\n"
"│  ├─ Dockerfile           # 单阶段构建\n"
"│  ├─ requirements.txt\n"
"│  └─ data/                # 运行时数据（DB + .secret_key），需持久化\n"
"├─ docker-compose.yml\n"
"├─ DEPLOY.md               # 离线部署说明\n"
"└─ backend/offline/        # 离线交付物与构建脚本\n"
"   ├─ build_image.sh       # 前端构建 + docker build + docker save\n"
"   ├─ gen_cert.sh          # 自签名证书生成\n"
"   └─ password_manager_image.tar  # 已导出镜像（约 216MB）"
)
note("offline/ 目录是存在的，包含 build_image.sh、gen_cert.sh 与导出的 tar 包，"
     "离线部署直接 load tar 即可，无需联网。", "info")

# ================================================================
# 5. 后端架构
# ================================================================
h1("5. 后端架构与代码详解")

# 5.1 main.py
h2("5.1 入口与装配 —— app/main.py")
para("服务在启动时通过 lifespan 调用 seed.seed() 完成建表、管理员、默认密钥与默认分组的初始化；"
     "这是自动迁移的触发点。CORS 仅允许显式配置的来源，且 allow_credentials=False（Bearer 无 Cookie）。")
para("静态资源必须用 NoCacheStaticFiles 子类托管，因为 @app.middleware 不会拦截已挂载的 StaticFiles 子应用；"
     "否则浏览器可能缓存旧 app.js，导致新 html 访问不到对应元素而崩溃。")
code_block(
"# main.py 路由挂载与静态托管（示意）\n"
"app.include_router(auth.router)\n"
"app.include_router(admin.mine_router)\n"
"app.include_router(admin.users_router)\n"
"app.include_router(admin.groups_router)\n"
"app.include_router(admin.audit_router)\n"
"app.include_router(passwords.router)\n"
"app.include_router(history.router)\n"
"app.include_router(keys.router)\n"
"app.include_router(orgkeys_router)\n"
"app.include_router(users_batch.router)\n"
"app.include_router(permissions.router)\n"
"app.mount('/', NoCacheStaticFiles(directory=str(STATIC_DIR), html=True), name='static')"
)
code_block(
"# 静态资源禁止缓存（防前端资源不一致）\n"
"class NoCacheStaticFiles(StaticFiles):\n"
"    async def get_response(self, path, scope):\n"
"        response = await super().get_response(path, scope)\n"
"        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'\n"
"        return response"
)

# 5.2 models
h2("5.2 数据模型 —— app/models.py + app/db.py")
para("db.py 关键设计：SQLite 引擎开启 check_same_thread=False 与 timeout=30，"
     "并通过事件监听器给每个连接执行 PRAGMA journal_mode=WAL 与 busy_timeout=30000，"
     "用“WAL + 忙等待”双保险避免 database is locked。")
code_block(
"# db.py —— 通用增量加列迁移（向后兼容核心）\n"
"def _migrate_columns():\n"
"    insp = inspect(engine)\n"
"    existing = {t: {c['name'] for c in insp.get_columns(t)}\n"
"                for t in insp.get_table_names()}\n"
"    for table_name, table in Base.metadata.tables.items():\n"
"        cols = existing.get(table_name)\n"
"        if cols is None:\n"
"            continue  # 整表缺失交给 create_all\n"
"        for col in table.columns:\n"
"            if col.name in cols:\n"
"                continue\n"
"            conn.execute(text(f'ALTER TABLE \"{table_name}\" '\n"
"                             f'ADD COLUMN \"{col.name}\" {column_decl(col)}'))"
)
note("新增数据库列：只要改 models.py 给模型加字段，_migrate_columns 会在启动时自动 ALTER TABLE 补齐，"
     "绝不要手写“给某表加某列”的硬编码迁移（历史上曾漏掉 passwords.deleted 导致 500）。", "info")
para("核心模型（表名 / 关键字段 / 说明）：")
table(
    ["模型", "表", "关键字段", "说明"],
    [
        ["User", "users", "username(唯一), hashed_password, pw_salt, pw_verifier, is_admin", "用户；pw_verifier=SM3(password||salt)"],
        ["Group", "groups", "name(唯一), description", "分组（多租户隔离单元）"],
        ["KeyRecord", "keys", "algorithm(gpg|sm2 唯一), public_key, private_key", "服务端默认密钥对"],
        ["OrgKey", "org_keys", "group_id, algorithm, public/private_key, fingerprint, private_protected, private_passphrase", "组织级密钥库"],
        ["PasswordEntry", "passwords", "title, username, algorithm, scheme, ciphertext, entry_salt, entry_iv, group_id, orgkey_id, deleted", "密码条目（软删）"],
        ["UserPermission", "user_permissions", "user_id(PK), perms(JSON)", "逐用户授权（授权即限制）"],
        ["History", "history", "password_id, group_id, action, ciphertext, comment", "审计日志（非 AuditLog）"],
        ["AuthSession", "auth_sessions", "jti(PK), user_id, last_activity, revoked, ip", "服务端会话，支持强制失效"],
    ],
    widths=[1.0, 0.9, 2.6, 2.0],
)
para("两个关联表（定义在 models.py 顶层，非类）：")
code_block(
"user_groups = Table('user_groups', Base.metadata,        # 用户 <-> 分组 多对多\n"
"    Column('user_id', ForeignKey('users.id', ondelete='CASCADE'), primary_key=True),\n"
"    Column('group_id', ForeignKey('groups.id', ondelete='CASCADE'), primary_key=True))\n"
"user_admin_groups = Table('user_admin_groups', ...)      # 分组管理员管理范围"
)
note("History 模型（不是 AuditLog）就是审计日志；前端“历史/审计”页的数据来自它。", "warn")

# 5.3 认证
h2("5.3 认证系统 —— routers/auth.py + security.py")
para("SCRAM-SM3 登录协议：")
bullet("login/begin：服务端返回用户库中持久化的 salt（非随机）+ 一次性 nonce + iter。")
bullet("客户端算 T = SM3(password || salt)，proof = SM3(T || nonce)。")
bullet("login/verify：服务端 expected = SM3(pw_verifier || nonce)，用 compare_digest 恒定时间比对。")
bullet("nonce 单次有效（内存 OrderedDict 存储），消费即删，防重放。")
code_block(
"# security.py:26  JWT 签发（HS256 + jti）\n"
"def create_token(username, jti):\n"
"    expire = datetime.now(timezone.utc) + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)\n"
"    payload = {'sub': username, 'exp': expire, 'jti': jti}\n"
"    return jwt.encode(payload, SECRET_KEY, algorithm='HS256')"
)
code_block(
"# security.py   verifier 派生与期望证明\n"
"def derive_password_verifier(password, salt_hex):\n"
"    return _sm3_hex((password or '').encode('utf-8') + bytes.fromhex(salt_hex or '00'))\n"
"def expected_proof(pw_verifier_hex, nonce_hex):\n"
"    msg = bytes.fromhex(pw_verifier_hex or '') + bytes.fromhex(nonce_hex or '')\n"
"    return _sm3_hex(msg)\n"
"def consume_login_challenge(username, nonce_hex):  # 一次性消费\n"
"    ch = _login_challenges.pop(username, None)\n"
"    if ch is None: return False\n"
"    if time.time() > ch['expires']: return False\n"
"    return secrets.compare_digest(ch['nonce'], nonce_hex or '')"
)
para("auth.py 关键路由与限速：")
code_block(
"# auth.py  登录限速（内存固定窗口 20/60s）\n"
"def _login_rate_limit(request):\n"
"    client = request.client.host if request.client else 'unknown'\n"
"    now = time.time()\n"
"    hits = _login_hits.setdefault(client, [])\n"
"    _login_hits[client] = [t for t in hits if now - t < _LOGIN_WINDOW]\n"
"    if len(_login_hits[client]) >= _LOGIN_LIMIT:\n"
"        raise HTTPException(429, detail='登录尝试过于频繁，请稍后再试')\n"
"    _login_hits[client].append(now)"
)
bullet("login/verify：先 consume_login_challenge 校验 nonce 未被消费/过期 → expected_proof 比对 → 签发 JWT。")
bullet("change-password/*：复用 SCRAM 校验当前密码（零明文），同时更新 pw_salt/pw_verifier 与 hashed_password；legacy 账号可用 current_password 明文兜底。")
bullet("单账号单会话：_enforce_single_session 会吊销其它 IP 的登录态（revoke_other_sessions + create_session）。")
note("登录挑战 _login_challenges 是进程内内存字典，仅单进程 uvicorn 成立。若用多 worker 部署，"
     "SCRAM 登录会因消费不到挑战而失败。生产请用单 worker 或改用共享存储。", "warn")
note("限速键用 request.client.host，而审计会话 IP 用 X-Forwarded-For。反代未把真实 IP 传给 client.host 时，"
     "限速会按反代 IP 聚合（所有用户共用一个桶）。", "warn")
note("security.py 中 derive_password_verifier 在第 69 行与第 78 行被定义了两次（内容相同，后者覆盖前者）。"
     "无害但属代码气味，建议删除其一。", "warn")

# 5.4 加密体系
h2("5.4 加密体系 —— app/crypto/")
h3("5.4.1 内层 SM4-CBC（entry_cipher.py，零知识）")
para("密钥派生用 PBKDF2-SM3（自实现迭代 XOR）；加密时随机 salt+iv，明文前加 MAGIC 魔数 b'PWM1' → SM4-CBC；"
     "密文/盐/iv 全以 hex 存库，服务端不持有条目密码。解密后校验 MAGIC，不匹配即 WrongPasswordError（密码错）。")
code_block(
"# entry_cipher.py  内层加密\n"
"def encrypt_entry(plaintext, password) -> dict:\n"
"    salt = os.urandom(16)\n"
"    key = _pbkdf2_sm3(password, salt)\n"
"    iv = os.urandom(16)\n"
"    pt = _pkcs7_pad(MAGIC + plaintext.encode('utf-8'))\n"
"    crypt = sm4.CryptSM4(); crypt.set_key(key, sm4.SM4_ENCRYPT)\n"
"    ct = crypt.crypt_cbc(iv, pt)\n"
"    return {'salt': salt.hex(), 'iv': iv.hex(),\n"
"            'ciphertext': ct.hex() if isinstance(ct, bytes) else bytes(ct).hex()}"
)
note("entry_cipher.py:32-33 的 fallback 分支 _sm3_hex 返回的是 (hash, sm4) 元组而非字符串。"
     "因 cryptography 必装、try 分支始终生效，此 fallback 当前不触发；但若将来 cryptography 缺失，SM3 会直接报错。建议修成只返回字符串。", "warn")
h3("5.4.2 OrgKey/服务端密钥调度（manager.py）")
bullet("PROVIDERS = {'gpg':..., 'sm2':...}，按算法名分发到 gpg_crypto / sm2_crypto。")
bullet("ensure_keys 保证两套默认密钥对存在（缺失自动生成）。")
bullet("decrypt_with_orgkey 解密时若 OrgKey 私钥受保护，自动用 OrgKey.private_passphrase 解锁。")
h3("5.4.3 gpg_crypto.py / sm2_crypto.py")
bullet("GPG(pgpy)：导入前先打 imghdr 兼容垫片（Python 3.13 已移除 imghdr）；遍历主密钥+子密钥解密，受口令私钥用 with k.unlock(passphrase)。")
bullet("SM2(gmssl)：generate_keypair 由私钥推公钥；decrypt 的 passphrase 形参对 SM2 无效（仅兼容接口，被忽略）。")

# 5.5 核心路由 passwords
h2("5.5 核心路由 —— routers/passwords.py")
para("所有接口挂 Depends(get_current_user)；写操作再叠加 require_perm。")
h3("序列化与 needs_password 判定")
code_block(
"# passwords.py  _serialize_meta（列表/详情统一序列化）\n"
"def _serialize_meta(db, e):\n"
"    return {\n"
"        'id': e.id, 'title': e.title or '', 'system': e.system or '',\n"
"        'username': e.username, 'algorithm': e.algorithm, 'scheme': e.scheme,\n"
"        'needs_password': bool(e.entry_salt) and bool(e.entry_iv),  # 有盐+iv 即需条目密码\n"
"        'notes': e.notes, 'group_id': e.group_id, 'orgkey_id': e.orgkey_id,\n"
"        'key_name': key_name, 'created_at': ..., 'updated_at': ...,\n"
"        'created_by': e.created_by, 'updated_by': e.updated_by,\n"
"    }"
)
h3("创建（409 去重）")
para("同一分组下，按 title（回退 username）归一化小写 + algorithm 完全相同则拒绝新增。")
code_block(
"# passwords.py  创建去重（409）\n"
"existing = db.query(PasswordEntry).filter_by(\n"
"    group_id=req.group_id, algorithm=algo, deleted=False).all()\n"
"for r in existing:\n"
"    if (r.title or r.username or '').strip().lower() == t_name:\n"
"        raise HTTPException(409, detail='该分组下已存在...')"
)
h3("导入 / 模板 / 导出契约")
bullet("POST /api/passwords/import：multipart/form-data，必填 group_id(int)、algorithm、entry_password、file(.xlsx)。逐行创建，部分失败不中断，回执含 total/created/skipped/errored。")
bullet("GET /api/passwords/template?fmt=xlsx：仅 xlsx 可用（fmt=csv 返回 400）。表头别名兼容新模板（密码文件名称/系统/用户名）与旧导出（标题/账号）。")
bullet("导出：plaintext=False 仅含密文；plaintext=True 需 {id:解密密码}，任一条解密失败整体拒绝（避免残缺明文）。")
h3("解密闭包 unlock 与 编辑/删除")
bullet("POST /api/passwords/{pid}/unlock：用请求体 JSON 传 entry_password（避免出现在 URL/日志）。")
bullet("_decrypt_entry_secret：无内层旧 legacy → 直接服务端密钥解；scheme='entry' → 直接 SM4 解；scheme='legacy' → 先 OrgKey/服务端私钥解外层、再 SM4 解内层。")
bullet("PUT /{pid} 编辑去重：与 create 一致（排除自身）；secret 为空串视为“保持不变”（编辑标题时不误清空密码）。")
bullet("DELETE /{pid}：软删除（置 deleted=True）。")

# 5.6 管理与多租户
h2("5.6 管理与多租户隔离 —— routers/admin.py + core/deps.py")
para("管理员双形态：is_admin=True 且 user_admin_groups 为空 = 超级管理员（见全部分组/用户/审计）；"
     "user_admin_groups 非空 = 分组管理员（仅管管理范围内的分组、新建分组自动纳入其范围、不能改其他管理员）。")
code_block(
"# admin.py  分组管理员可见用户 = 与自己管理分组有交集者（含自己）\n"
"def _visible_user_ids(db, caller):\n"
"    if is_global_admin(db, caller): return None\n"
"    my_admin_ids = get_admin_group_ids(db, caller)\n"
"    ...\n"
"    for u in rows:\n"
"        u_ids = {g.id for g in u.groups} | get_admin_group_ids(db, u)\n"
"        if u.id == caller.id or (u_ids & my_admin_ids): out.add(u.id)\n"
"    return out"
)
para("多租户隔离闸门集中在 core/deps.py：")
table(
    ["闸门", "位置", "作用"],
    [
        ["visibility_filter(column, user, group_ids)", "deps.py:122", "列表查询注入 group_id IN (...)；管理员返回 None（不过滤）"],
        ["ensure_group_access(db, user, group_id)", "deps.py:133", "写入/读单条前校验分组归属（403）"],
        ["get_user_group_ids / get_user_groups", "deps.py:88/118", "计算可见分组（普通=所属；分组管理员=所属∪管理；超级=全部）"],
        ["is_global_admin", "deps.py:83", "超级管理员判定"],
        ["_visible_user_ids", "admin.py:120", "管理员视角下用户/分组/审计可见范围"],
        ["require_perm(key)", "perms.py:85", "逐用户授权（404=全开，有记录=白名单）"],
        ["get_current_user → is_session_valid", "deps.py / sessions.py", "令牌 + 服务端会话双校验，吊销/空闲即 401"],
    ],
    widths=[2.4, 1.3, 3.1],
)
note("改任何数据可见性（密码/密钥/导出/审计）都只需改 deps.get_user_groups / visibility_filter —— 它们是统一闸门，"
     "不要在各路由里各自判断分组。", "info")

# 5.7 其它路由
h2("5.7 其余路由")
bullet("routers/keys.py + orgkeys_router：/api/keys/status 就绪检查；OrgKey 的 generate/import/export/batch-delete/delete；导出需管理员 + ensure_group_access。")
bullet("routers/users_batch.py：/api/admin/users/template + /batch，xlsx 批量建用户。")
bullet("routers/history.py：GET /api/passwords/{pid}/history，先 ensure_group_access 防越权。")
bullet("routers/permissions.py：GET/PUT/DELETE /api/admin/permissions/users/{uid}，授权即限制。")

# 5.8 配置
h2("5.8 配置 —— app/config.py")
code_block(
"# config.py  JWT 密钥持久化（同一 data 卷 → 升级后旧令牌不失效）\n"
"SECRET_KEY_FILE = DATA_DIR / '.secret_key'\n"
"SECRET_KEY = os.getenv('SECRET_KEY')\n"
"if not SECRET_KEY:\n"
"    if SECRET_KEY_FILE.exists(): SECRET_KEY = SECRET_KEY_FILE.read_text().strip()\n"
"    else: SECRET_KEY = secrets.token_hex(32); SECRET_KEY_FILE.write_text(SECRET_KEY)"
)
para("关键配置：DATA_DIR(backend/data)、DB_PATH、SECRET_KEY（环境变量或持久化）、ALGORITHM='HS256'、"
     "ACCESS_TOKEN_EXPIRE_MINUTES(默认1440)、SESSION_IDLE_SECONDS(默认600)、ADMIN_USERNAME/PASSWORD、"
     "ALLOW_REGISTRATION、REGISTER_DEFAULT_GROUP、HOST/PORT。")

# ================================================================
# 6. 前端
# ================================================================
h1("6. 前端架构与代码详解")
h2("6.1 入口 App.vue / main.js")
para("main.js 极简：createApp(App).mount('#app')。App.vue 用 loggedIn 切换 Login / 主界面；"
     "进入应用调用 enterApp()，三类数据并行预加载，避免首屏点“查看”时 state.entries 未填充。")
code_block(
"// App.vue  enterApp —— 并行预加载消除首屏竞态\n"
"async function enterApp() {\n"
"  await Promise.all([\n"
"    loadKeysStatus(),\n"
"    loadEntries(),\n"
"    loadOrgKeys(),\n"
"  ])\n"
"}"
)
note("此前“首次登录点击查看显示全 —”的 Bug 根因，正是 enterApp 串行 await（先 loadKeysStatus 再 loadEntries），"
     "ViewModal 依赖 state.entries 查找条目导致字段全空。改用 Promise.all 后已修复。", "info")

h2("6.2 状态 store.js")
code_block(
"// store.js  权限判断：null=全部可用；数组=白名单\n"
"export function can(key) {\n"
"  const p = state.permissions\n"
"  if (p === null || p === undefined) return true\n"
"  return p.includes(key)\n"
"}"
)
bullet("state.permissions 来自 /me：null = 管理员或未被授权过（全开）；数组 = 仅清单内可用。")
bullet("空闲超时：IDLE_TIMEOUT_MS=600000（与服务端 600s 对齐）；onActivity 监听全局事件刷新倒计时，并每 ~20s 上报心跳 apiActivity()，防止纯前端操作不打 API 导致令牌被服务端吊销。")

h2("6.3 API 层 http.js + auth.js")
code_block(
"// http.js  错误透传：e.message = detail/message；任意 401 触发登出\n"
"if (!res.ok) {\n"
"  const msg = (data && (data.detail || data.message)) || '请求失败 (' + res.status + ')'\n"
"  const e = new Error(typeof msg === 'string' ? msg : JSON.stringify(msg))\n"
"  e.status = res.status\n"
"  if (res.status === 401) fireUnauthorized(msg)\n"
"  throw e\n"
"}"
)
note("上传文件走 FormData 时不要手动加 Content-Type: application/json（http.js 已对 FormData 跳过 JSON 头，"
     "后端 python-multipart 需要 multipart 边界）。", "info")

h2("6.4 关键组件")
table(
    ["组件", "作用与数据流"],
    [
        ["PasswordPanel.vue", "列表/分页/搜索/筛选；打开查看→ViewModal，编辑→PasswordFormModal。分页接口拿 items/total，state.entries 保持全量供查重/导出。"],
        ["ViewModal.vue", "props entry（非 id）；onMounted 若 needs_password 则锁住，否则 POST /{pid}/unlock 取明文。解锁后显示并支持复制。"],
        ["PasswordFormModal.vue", "新增/编辑表单；前端去重预校验（与后端一致：title+algorithm 同分组、排除自身）；OrgKey 选项按 algorithm+groupId 动态拉取。"],
        ["Login.vue", "先 registerStatus() 决定是否展现注册入口；登录调 doLogin→SCRAM；409 回退 /api/auth/login 一次性迁移。"],
        ["ChangePwModal.vue", "changePassword(current,next)，SCRAM 优先、legacy 兜底。"],
        ["KeyPanel.vue / KeyGenModal / KeyImportModal", "密钥库分页、生成/导入/导出（公钥必显、私钥需 has_private）、批量删除。"],
        ["AdminModal / UserFormModal / GroupFormModal / PermPanel / HistoryModal / ImportModal / ExportModal", "管理后台：用户/分组/授权/审计/导入导出。"],
    ],
    widths=[2.0, 4.8],
)
code_block(
"// PasswordFormModal.vue  前端去重预校验（与后端一致）\n"
"const dup = state.entries.find(\n"
"  (e) => e.group_id === gid &&\n"
"         (e.title || e.username || '').trim().toLowerCase() === tNorm &&\n"
"         e.algorithm === algo)\n"
"if (dup) return (formError.value = '该分组下已存在...请勿重复新增')"
)

h2("6.5 前端 SM3 计算 crypto/sm3.js")
para("前端纯 JS 实现 GM/T 0003-2012 SM3，必须与后端 cryptography 的 SM3 逐字节一致（已统一算法保证）。")
code_block(
"// crypto/sm3.js  T = SM3(password || salt); proof = SM3(T || nonce)\n"
"export function scramProof(password, saltHex, nonceHex) {\n"
"  const saltBytes = hexToBytes(saltHex), nonceBytes = hexToBytes(nonceHex)\n"
"  const pwBytes = new TextEncoder().encode(password)\n"
"  const tInput = new Uint8Array(pwBytes.length + saltBytes.length)\n"
"  tInput.set(pwBytes, 0); tInput.set(saltBytes, pwBytes.length)\n"
"  const verifier = sm3Bytes(tInput)               // T\n"
"  const proofInput = new Uint8Array(verifier.length + nonceBytes.length)\n"
"  proofInput.set(verifier, 0); proofInput.set(nonceBytes, verifier.length)\n"
"  return bytesToHex(sm3Bytes(proofInput))         // proof\n"
"}"
)

# ================================================================
# 7. 关键业务流程
# ================================================================
h1("7. 关键业务流程")
h2("7.1 登录（SCRAM-SM3）")
code_block(
"1. 前端 POST /api/auth/login/begin {username}\n"
"2. 后端返回持久化 pw_salt + 一次性 nonce\n"
"3. 前端算 T=SM3(password||salt_raw), proof=SM3(T||nonce_raw)\n"
"4. 前端 POST /api/auth/login/verify {username,nonce,proof}\n"
"5. 后端 expected=SM3(pw_verifier||nonce)，compare_digest 比对 → 签发 JWT\n"
"6. 旧用户首次走 /api/auth/login 明文登录成功会自动迁移为 SCRAM"
)
h2("7.2 创建密码（双加密层）")
bullet("symmetric：仅内层 SM4（scheme='entry'）。")
bullet("gpg/sm2：内层 SM4 后再用 OrgKey 公钥（或服务端密钥）做外层非对称（scheme='legacy'）；查看需“私钥 + 条目密码”两层。")
h2("7.3 查看密码（unlock）")
bullet("前端打开 ViewModal，传入整条 entry 对象。")
bullet("needs_password 为 true → 弹窗要求输入 entry_password，POST /{pid}/unlock {entry_password}。")
bullet("needs_password 为 false → 直接 POST /{pid}/unlock {} 取明文（legacy 纯密文条目无条目密码层）。")
h2("7.4 管理员双形态")
bullet("超级管理员：is_admin=True 且 user_admin_groups 空 → 全量可见。")
bullet("分组管理员：仅见“所属分组 ∪ 管理的分组”，不能改其他管理员、不能删超级管理员/自己；新建分组自动纳入其范围。")
bullet("/api/auth/me 返回 is_global_admin 供前端显隐管理员开关。")

# ================================================================
# 8. 构建与部署
# ================================================================
h1("8. 构建与部署")
h2("8.1 离线交付物")
bullet("backend/offline/password_manager_image.tar（约 216MB）：已导出的镜像，docker load 即用。")
bullet("backend/offline/build_image.sh：前端 npm build + docker build --platform linux/amd64 + docker save。")
bullet("backend/offline/gen_cert.sh：自签名证书生成（供 HTTPS 部署）。")
bullet("docker-compose.yml / DEPLOY.md：离线部署说明与变量表。")
h2("8.2 镜像构建要点（Dockerfile 单阶段）")
bullet("FROM python:3.13-slim；先装 requirements（默认 PyPI，OFFLINE=1 时 --no-index 用 offline/wheels）。")
bullet("COPY frontend/dist ./app/static（前端预编译产物直接拷入，容器内不含 Node）。")
bullet("CMD [\"python\",\"run.py\"]，VOLUME [\"/app/data\"]，暴露 9010/9080。")
note("Dockerfile 静态资源 COPY 路径必须写成 ./app/static（→ 容器内 /app/app/static，正是 main.py 的 STATIC_DIR），"
     "绝不可写 ./app/app/static（会落 /app/app/app/static，应用读不到）。", "warn")
h2("8.3 运行入口 run.py（HTTPS）")
bullet("设了 SSL_CERTFILE/SSL_KEYFILE 即启用 TLS（uvicorn ssl_certfile/ssl_keyfile）。")
bullet("SSL_REDIRECT!=0 时另起线程把明文 9080 全部 307 跳到 HTTPS，跳转 Host 取自请求头（不硬编码）。")
h2("8.4 部署与升级")
code_block(
"# 离线部署\n"
"docker load -i backend/offline/password_manager_image.tar\n"
"docker compose up -d\n"
"# 升级：保留旧 backend/data/ 目录即可不丢数据（DB + .secret_key 都在其中）"
)
note("升级时 data/ 卷必须连同 *.db-wal / *.db-shm 一起保留（WAL 伴随文件），否则可能丢最近写入。", "warn")
note("默认账号 admin/admin123 是弱口令，上线第一要务是改强口令或启用 HTTPS。.env 含机密不可提交。", "warn")

# ================================================================
# 9. 维护要点与常见坑
# ================================================================
h1("9. 维护要点与常见坑")
h2("9.1 必看：多租户隔离闸门")
para("任何“谁能看什么数据”的问题，先看 deps.py 的 visibility_filter / ensure_group_access / get_user_group_ids，"
     "以及 admin.py 的 _visible_user_ids。改可见性只动这些地方，不要散落到各路由。")
h2("9.2 易踩的坑清单")
table(
    ["#", "坑", "后果 / 处理"],
    [
        ["1", "offline/ 脚本缺失（误传）", "实际 build_image.sh / gen_cert.sh / tar 均存在，离线部署照常。"],
        ["2", "登录挑战是进程内内存字典", "仅单进程 uvicorn 成立；多 worker 会导致 SCRAM 登录失败。"],
        ["3", "限速键来源不一致", "client.host vs X-Forwarded-For；反代未透传真实 IP 时限速聚合到反代 IP。"],
        ["4", "derive_password_verifier 重复定义", "security.py:69 与 :78，无害但属代码气味，建议删除其一。"],
        ["5", "entry_cipher.py:32-33 fallback 返回元组", "当前因 cryptography 必装不触发，但属潜在隐患，建议修成只返回字符串。"],
        ["6", "编辑去重以 title||username 归一化", "title 可空时以 username 去重；两者皆空仍可创建。"],
        ["7", "前端分页 vs 全量不一致风险", "列表用分页接口，查重/导出用 state.entries 全量；afterSaved 会同时刷新。"],
        ["8", "legacy 纯密文无条目密码层", "旧数据 entry_salt/entry_iv 为空，通过分组校验的人都能看到明文。"],
        ["9", "admin123 默认弱口令", "未配 .env 即明文 HTTP 运行，上线必改强口令或启用 HTTPS。"],
        ["10", "静态资源必须 No-Cache", "否则旧 app.js + 新 html 可能导致 handler 崩溃（已用 NoCacheStaticFiles 解决）。"],
        ["11", "WAL 伴随文件需一起保留", "升级别只拷 .db，*.db-wal/*.db-shm 也要保留，否则丢最近写入。"],
        ["12", "Dockerfile 静态 COPY 路径", "必须 ./app/static，写 ./app/app/static 会读不到页面。"],
    ],
    widths=[0.4, 2.2, 4.2],
)

# ================================================================
# 10. 常见问题排查
# ================================================================
h1("10. 常见问题排查（FAQ）")
bullet("登录返回 409：前端回退明文 /login 一次性迁移；若是 SCRAM 阶段 409，检查 salt/nonce 是否由 begin 正确返回。")
bullet("查看密码显示全 —：检查 ViewModal 是否拿到真实 entry 对象（而非依赖未就绪的 state.entries）；确认 App.enterApp 已并行预加载。")
bullet("新建/导入报 409：同分组下“密码文件名称 + 加密方式”重复，改名或换算法。")
bullet("解密报错 WrongPasswordError：entry_password 错误，或 legacy 条目的服务端密钥不可用。")
bullet("列表为空但接口有数据：检查多租户闸门与 /me 的 is_global_admin、分组归属。")
bullet("升级后旧令牌失效：SECRET_KEY 未持久化到 data/.secret_key 导致每次随机，需保留 data/ 卷。")
bullet("Docker 构建后 GET / 404：前端产物未拷入 /app/app/static（commit 增量法漏拷 frontend/dist）。")

# ================================================================
# 11. 附录：关键接口速查
# ================================================================
h1("11. 附录：关键接口速查")
table(
    ["方法", "路径", "说明"],
    [
        ["POST", "/api/auth/login/begin", "取 salt+nonce（SCRAM 第一步）"],
        ["POST", "/api/auth/login/verify", "校验 proof，签发 JWT"],
        ["POST", "/api/auth/login", "legacy 明文登录（自动迁移）"],
        ["POST", "/api/auth/change-password/begin|verify", "自助改密"],
        ["GET", "/api/auth/me", "当前用户与权限、is_global_admin"],
        ["GET", "/api/passwords", "分页列表（items/total）"],
        ["POST", "/api/passwords", "创建（409 去重）"],
        ["POST", "/api/passwords/import", "批量导入（multipart）"],
        ["GET", "/api/passwords/template?fmt=xlsx", "下载模板"],
        ["POST", "/api/passwords/{pid}/unlock", "解密查看（JSON 传 entry_password）"],
        ["PUT", "/api/passwords/{pid}", "编辑（排除自身去重）"],
        ["DELETE", "/api/passwords/{pid}", "软删除"],
        ["GET", "/api/passwords/{pid}/history", "审计历史"],
        ["GET", "/api/admin/groups", "分组列表（注意是 /api/admin/ 前缀）"],
        ["GET/PUT/DELETE", "/api/admin/permissions/users/{uid}", "逐用户授权"],
    ],
    widths=[1.0, 3.4, 2.4],
)
note("路由顺序陷阱：Starlette 中 {pid}:int 先匹配路径再校验 int。template/import 等非 int 段必须定义在 "
     "@router.get('/{pid}') 之前，否则会 422。", "warn")

doc.save(OUT)
print("SAVED:", OUT)
